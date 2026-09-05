"""Provider-neutral page extraction for smart-import source documents."""
from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import shutil
import subprocess
import tempfile
from typing import Any, BinaryIO, Protocol
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.oxml.ns import qn
from PIL import Image
from pypdf import PdfReader
from openpyxl import load_workbook


MAX_DOCUMENT_PAGES = 500
MAX_PAGE_TEXT_CHARS = 200_000
MIN_MEANINGFUL_TEXT_CHARS = 20
MAX_DOCX_ENTRIES = 5_000
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 200
MAX_IMAGE_PIXELS = 50_000_000
MAX_XLSX_ROWS_PER_SHEET = 100_000
MAX_XLSX_COLUMNS = 1_000
MAX_LEGACY_DOC_TEXT_BYTES = 20 * 1024 * 1024


class DocumentParseError(ValueError):
    """A safe, user-facing document parsing error."""


@dataclass(frozen=True)
class ParsedPage:
    page_number: int
    text_content: str
    ocr_status: str
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ParsedDocument:
    pages: list[ParsedPage]
    parser: str
    page_numbering: str = "physical"


class DocumentParser(Protocol):
    def parse(
        self, stream: BinaryIO, original_filename: str, mime_type: str | None = None
    ) -> ParsedDocument:
        ...


class DefaultDocumentParser:
    """Extract embedded text and identify pages that need a later OCR pass."""

    def parse(
        self, stream: BinaryIO, original_filename: str, mime_type: str | None = None
    ) -> ParsedDocument:
        del mime_type
        suffix = Path(original_filename).suffix.lower()
        try:
            if suffix == ".pdf":
                return self._parse_pdf(stream)
            if suffix in {".dxf", ".dwg"}:
                from app.services.cad_conversion_service import cad_to_pdf
                parsed = self._parse_pdf(BytesIO(cad_to_pdf(stream, original_filename)))
                return ParsedDocument(parsed.pages, "cad_to_pdf")
            if suffix == ".docx":
                return self._parse_docx(stream)
            if suffix == ".doc":
                return self._parse_doc(stream)
            if suffix == ".xlsx":
                return self._parse_xlsx(stream)
            if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
                return self._parse_image(stream, suffix)
            raise DocumentParseError("该文件类型暂不支持分页解析")
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError("文档损坏、加密或格式无法识别") from exc

    def _parse_pdf(self, stream: BinaryIO) -> ParsedDocument:
        reader = PdfReader(stream, strict=False)
        if reader.is_encrypted:
            try:
                unlocked = reader.decrypt("")
            except Exception as exc:
                raise DocumentParseError("加密 PDF 无法解析，请先移除密码") from exc
            if not unlocked:
                raise DocumentParseError("加密 PDF 无法解析，请先移除密码")
        if len(reader.pages) > MAX_DOCUMENT_PAGES:
            raise DocumentParseError(f"文档页数不能超过 {MAX_DOCUMENT_PAGES} 页")

        pages: list[ParsedPage] = []
        for number, page in enumerate(reader.pages, start=1):
            text = _clean_text(page.extract_text() or "")
            has_images = _pdf_page_has_images(page)
            meaningful_chars = len("".join(text.split()))
            # Scans can contain a searchable header while the actual form is
            # an image; CAD exports can contain only vector outlines.
            vector_candidate = meaningful_chars < 200 and _pdf_vector_outline_candidate(page)
            needs_ocr = has_images or vector_candidate or meaningful_chars < MIN_MEANINGFUL_TEXT_CHARS
            pages.append(
                ParsedPage(
                    page_number=number,
                    text_content=_bounded_text(text),
                    ocr_status="pending" if needs_ocr else "not_required",
                    metadata={
                        "source_format": "pdf",
                        "text_chars": len(text),
                        "has_images": has_images,
                        "vector_outline_candidate": vector_candidate,
                        "is_scanned_candidate": needs_ocr,
                        "width_points": _as_float(page.mediabox.width),
                        "height_points": _as_float(page.mediabox.height),
                    },
                )
            )
        return ParsedDocument(pages=pages, parser="pypdf")

    def _parse_docx(self, stream: BinaryIO) -> ParsedDocument:
        _validate_docx_archive(stream)
        document = Document(stream)
        logical_pages: list[list[str]] = [[]]

        for block in document.element.body.iterchildren():
            if block.tag == qn("w:p"):
                segments = _docx_paragraph_segments(block)
                for index, segment in enumerate(segments):
                    if segment:
                        logical_pages[-1].append(segment)
                    if index < len(segments) - 1:
                        logical_pages.append([])
            elif block.tag == qn("w:tbl"):
                table_text = _docx_table_text(block)
                if table_text:
                    logical_pages[-1].append(table_text)

        if len(logical_pages) > MAX_DOCUMENT_PAGES:
            raise DocumentParseError(f"文档逻辑页数不能超过 {MAX_DOCUMENT_PAGES} 页")
        pages = [
            ParsedPage(
                page_number=number,
                text_content=_bounded_text(_clean_text("\n".join(parts))),
                ocr_status="not_required",
                metadata={
                    "source_format": "docx",
                    "page_numbering": "logical",
                    "text_chars": len(_clean_text("\n".join(parts))),
                },
            )
            for number, parts in enumerate(logical_pages, start=1)
        ]
        return ParsedDocument(
            pages=pages, parser="python-docx", page_numbering="logical"
        )

    def _parse_doc(self, stream: BinaryIO) -> ParsedDocument:
        text = _extract_legacy_doc_text(stream)
        logical_pages = text.split("\f")
        if len(logical_pages) > MAX_DOCUMENT_PAGES:
            raise DocumentParseError(f"文档逻辑页数不能超过 {MAX_DOCUMENT_PAGES} 页")
        pages = [
            ParsedPage(
                page_number=number,
                text_content=_bounded_text(_clean_text(page)),
                ocr_status="not_required",
                metadata={
                    "source_format": "doc",
                    "page_numbering": "logical",
                    "text_chars": len(_clean_text(page)),
                },
            )
            for number, page in enumerate(logical_pages, start=1)
        ]
        if not any(page.text_content for page in pages):
            raise DocumentParseError("DOC 中没有可提取的文本，请转换为 DOCX 或 PDF")
        return ParsedDocument(
            pages=pages, parser="antiword", page_numbering="logical"
        )

    def _parse_image(self, stream: BinaryIO, suffix: str) -> ParsedDocument:
        with Image.open(stream) as image:
            frame_count = getattr(image, "n_frames", 1)
            if frame_count > MAX_DOCUMENT_PAGES:
                raise DocumentParseError(f"图像页数不能超过 {MAX_DOCUMENT_PAGES} 页")
            pages: list[ParsedPage] = []
            for number in range(1, frame_count + 1):
                image.seek(number - 1)
                if image.width * image.height > MAX_IMAGE_PIXELS:
                    raise DocumentParseError("扫描图片像素尺寸超过系统安全限制")
                pages.append(
                    ParsedPage(
                        page_number=number,
                        text_content="",
                        ocr_status="pending",
                        metadata={
                            "source_format": suffix.lstrip("."),
                            "has_images": True,
                            "is_scanned_candidate": True,
                            "width_pixels": image.width,
                            "height_pixels": image.height,
                        },
                    )
                )
        return ParsedDocument(pages=pages, parser="pillow")

    def _parse_xlsx(self, stream: BinaryIO) -> ParsedDocument:
        """Expose each worksheet as a logical page while preserving row boundaries."""
        _validate_xlsx_archive(stream)
        workbook = load_workbook(stream, read_only=True, data_only=True)
        if len(workbook.worksheets) > MAX_DOCUMENT_PAGES:
            raise DocumentParseError(f"Excel 工作表数不能超过 {MAX_DOCUMENT_PAGES} 个")
        pages: list[ParsedPage] = []
        try:
            for number, sheet in enumerate(workbook.worksheets, start=1):
                if (
                    sheet.max_row > MAX_XLSX_ROWS_PER_SHEET
                    or sheet.max_column > MAX_XLSX_COLUMNS
                ):
                    raise DocumentParseError("Excel 工作表行列数超过系统安全限制")
                rows: list[str] = []
                row_count = 0
                for values in sheet.iter_rows(values_only=True):
                    cells = [
                        "" if value is None else str(value).strip() for value in values
                    ]
                    while cells and not cells[-1]:
                        cells.pop()
                    if not any(cells):
                        continue
                    row_count += 1
                    rows.append("\t".join(cells))
                text = _bounded_text(_clean_text("\n".join(rows)))
                pages.append(
                    ParsedPage(
                        page_number=number,
                        text_content=text,
                        ocr_status="not_required",
                        metadata={
                            "source_format": "xlsx",
                            "page_numbering": "worksheet",
                            "sheet_name": sheet.title,
                            "row_count": row_count,
                            "text_chars": len(text),
                        },
                    )
                )
        finally:
            workbook.close()
        return ParsedDocument(
            pages=pages, parser="openpyxl", page_numbering="worksheet"
        )


def _validate_docx_archive(stream: BinaryIO) -> None:
    try:
        stream.seek(0)
        with ZipFile(stream) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise DocumentParseError("DOCX 文件条目数量超过系统安全限制")
            total_size = sum(entry.file_size for entry in entries)
            compressed_size = sum(max(entry.compress_size, 1) for entry in entries)
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise DocumentParseError("DOCX 解压后大小超过系统安全限制")
            if total_size / max(compressed_size, 1) > MAX_DOCX_COMPRESSION_RATIO:
                raise DocumentParseError("DOCX 压缩比例异常，已拒绝解析")
    except BadZipFile as exc:
        raise DocumentParseError("DOCX 文件损坏或格式无法识别") from exc
    finally:
        stream.seek(0)


def _extract_legacy_doc_text(stream: BinaryIO) -> str:
    executable = shutil.which("antiword")
    if not executable:
        raise DocumentParseError("服务器 DOC 解析组件未安装，请转换为 DOCX 或 PDF")
    stream.seek(0)
    with tempfile.TemporaryDirectory(prefix="weld-doc-") as directory:
        input_path = Path(directory) / "source.doc"
        output_path = Path(directory) / "output.txt"
        with input_path.open("wb") as target:
            shutil.copyfileobj(stream, target, length=1024 * 1024)
        try:
            with output_path.open("wb") as output:
                result = subprocess.run(
                    [executable, "-m", "UTF-8.txt", str(input_path)],
                    stdin=subprocess.DEVNULL,
                    stdout=output,
                    stderr=subprocess.PIPE,
                    timeout=30,
                    check=False,
                )
        except subprocess.TimeoutExpired as exc:
            raise DocumentParseError("DOC 解析超时，请转换为 DOCX 或 PDF") from exc
        if result.returncode != 0:
            raise DocumentParseError("DOC 文件损坏、加密或格式无法识别")
        if output_path.stat().st_size > MAX_LEGACY_DOC_TEXT_BYTES:
            raise DocumentParseError("DOC 提取后的文本超过系统安全限制")
        try:
            return output_path.read_text(encoding="utf-8")
        except UnicodeDecodeError as exc:
            raise DocumentParseError("DOC 文本编码无法识别") from exc
        finally:
            stream.seek(0)


def _validate_xlsx_archive(stream: BinaryIO) -> None:
    try:
        stream.seek(0)
        with ZipFile(stream) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise DocumentParseError("Excel 文件条目数量超过系统安全限制")
            total_size = sum(entry.file_size for entry in entries)
            compressed_size = sum(max(entry.compress_size, 1) for entry in entries)
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise DocumentParseError("Excel 解压后大小超过系统安全限制")
            if total_size / max(compressed_size, 1) > MAX_DOCX_COMPRESSION_RATIO:
                raise DocumentParseError("Excel 压缩比例异常，已拒绝解析")
    except BadZipFile as exc:
        raise DocumentParseError("Excel 文件损坏或格式无法识别") from exc
    finally:
        stream.seek(0)


def _pdf_page_has_images(page: Any) -> bool:
    visited = set()
    def inspect(resources, depth=0):
        if depth > 12:
            return True  # An uninspectable nested form needs visual review.
        resources = resources.get_object() if hasattr(resources, "get_object") else resources
        xobjects = (resources or {}).get("/XObject") or {}
        xobjects = xobjects.get_object() if hasattr(xobjects, "get_object") else xobjects
        for value in xobjects.values():
            obj = value.get_object() if hasattr(value, "get_object") else value
            if id(obj) in visited:
                continue
            visited.add(id(obj))
            if obj.get("/Subtype") == "/Image":
                return True
            if obj.get("/Subtype") == "/Form" and inspect(obj.get("/Resources"), depth+1):
                return True
        return False
    try:
        if inspect(page.get("/Resources")):
            return True
        return len(page.images) > 0
    except Exception:
        return True


def _pdf_vector_outline_candidate(page: Any) -> bool:
    try:
        content = page.get_contents()
        if content is None:
            return False
        # Glyph outlines use many curves/path segments despite little searchable text.
        return sum(op in {b"c", b"v", b"y", b"l"} for _, op in content.operations) >= 80
    except Exception:
        return True


def _docx_paragraph_segments(paragraph: Any) -> list[str]:
    segments = [""]
    for element in paragraph.iter():
        if element.tag == qn("w:t") and element.text:
            segments[-1] += element.text
        elif element.tag == qn("w:tab"):
            segments[-1] += "\t"
        elif element.tag == qn("w:br"):
            if element.get(qn("w:type")) == "page":
                segments.append("")
            else:
                segments[-1] += "\n"
    return segments


def _docx_table_text(table: Any) -> str:
    rows: list[str] = []
    for row in table.iterchildren(qn("w:tr")):
        cells: list[str] = []
        for cell in row.iterchildren(qn("w:tc")):
            text = "".join(node.text or "" for node in cell.iter(qn("w:t")))
            cells.append(text.strip())
        rows.append("\t".join(cells).rstrip())
    return "\n".join(rows).strip()


def _clean_text(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\x00", "").splitlines()]
    return "\n".join(lines).strip()


def _bounded_text(text: str) -> str:
    if len(text) > MAX_PAGE_TEXT_CHARS:
        raise DocumentParseError("单页文本内容异常，超过系统安全限制")
    return text


def _as_float(value: Any) -> float | None:
    try:
        return round(float(value), 2)
    except (TypeError, ValueError):
        return None
