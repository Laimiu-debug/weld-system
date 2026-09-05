from io import BytesIO
from time import perf_counter

import pytest
from docx import Document
from PIL import Image
from openpyxl import Workbook
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from app.services.document_parser_service import (
    DefaultDocumentParser,
    DocumentParseError,
)
from app.services import document_parser_service
from app.schemas.smart_import import BatchAIExtractionRequest
from pydantic import ValidationError


def _text_pdf() -> BytesIO:
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.drawString(72, 760, "WPS No. WPS-001 welding procedure")
    pdf.showPage()
    pdf.drawString(72, 760, "PQR No. PQR-001 qualification record")
    pdf.save()
    stream.seek(0)
    return stream


def _scanned_pdf() -> BytesIO:
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.drawInlineImage(Image.new("RGB", (120, 80), "white"), 72, 650)
    pdf.save()
    stream.seek(0)
    return stream


def _many_page_pdf(page_count: int) -> BytesIO:
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    for index in range(page_count):
        pdf.drawString(72, 760, f"WPS page {index + 1}: Q345R, 12 mm")
        pdf.showPage()
    pdf.save()
    stream.seek(0)
    return stream


def test_pdf_embedded_text_is_split_by_physical_page() -> None:
    parsed = DefaultDocumentParser().parse(_text_pdf(), "WPS.pdf")

    assert parsed.parser == "pypdf"
    assert parsed.page_numbering == "physical"
    assert [page.page_number for page in parsed.pages] == [1, 2]
    assert "WPS-001" in parsed.pages[0].text_content
    assert "PQR-001" in parsed.pages[1].text_content
    assert all(page.ocr_status == "not_required" for page in parsed.pages)


def test_image_only_pdf_is_marked_for_later_ocr() -> None:
    parsed = DefaultDocumentParser().parse(_scanned_pdf(), "scan.pdf")

    assert len(parsed.pages) == 1
    assert parsed.pages[0].ocr_status == "pending"
    assert parsed.pages[0].metadata["has_images"] is True
    assert parsed.pages[0].metadata["is_scanned_candidate"] is True


def test_xobject_image_pdf_is_also_marked_for_later_ocr() -> None:
    image_stream = BytesIO()
    Image.new("RGB", (120, 80), "white").save(image_stream, format="PNG")
    image_stream.seek(0)
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.drawImage(ImageReader(image_stream), 72, 650, width=120, height=80)
    pdf.save()
    stream.seek(0)

    parsed = DefaultDocumentParser().parse(stream, "xobject-scan.pdf")

    assert parsed.pages[0].ocr_status == "pending"


def test_scan_with_searchable_header_still_reads_image_form() -> None:
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.drawString(72, 760, "Procedure Qualification Record PQR-001")
    pdf.drawInlineImage(Image.new("RGB", (120, 80), "white"), 72, 650)
    pdf.save()
    stream.seek(0)
    assert DefaultDocumentParser().parse(stream, "PQR.pdf").pages[0].ocr_status == "pending"


def test_vector_outlines_without_text_require_vision_ocr() -> None:
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.rect(72, 300, 300, 300)
    pdf.save()
    stream.seek(0)
    assert DefaultDocumentParser().parse(stream, "CAD-export.pdf").pages[0].ocr_status == "pending"


def test_docx_uses_explicit_page_breaks_and_keeps_table_text() -> None:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("WPS No. WPS-002")
    document.add_page_break()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Base metal"
    table.cell(0, 1).text = "Q345R"
    document.save(stream)
    stream.seek(0)

    parsed = DefaultDocumentParser().parse(stream, "WPS.docx")

    assert parsed.page_numbering == "logical"
    assert len(parsed.pages) == 2
    assert "WPS-002" in parsed.pages[0].text_content
    assert "Base metal\tQ345R" in parsed.pages[1].text_content
    assert all(page.ocr_status == "not_required" for page in parsed.pages)


def test_multipage_tiff_creates_one_ocr_page_per_frame() -> None:
    stream = BytesIO()
    first = Image.new("RGB", (20, 30), "white")
    second = Image.new("RGB", (40, 50), "white")
    first.save(stream, format="TIFF", save_all=True, append_images=[second])
    stream.seek(0)

    parsed = DefaultDocumentParser().parse(stream, "welder-card.tiff")

    assert len(parsed.pages) == 2
    assert [page.ocr_status for page in parsed.pages] == ["pending", "pending"]
    assert parsed.pages[1].metadata["width_pixels"] == 40


def test_legacy_doc_uses_antiword_and_preserves_logical_pages(monkeypatch) -> None:
    monkeypatch.setattr(
        document_parser_service,
        "_extract_legacy_doc_text",
        lambda stream: "WPS No. WPS-003\fPQR No. PQR-003",
    )

    parsed = DefaultDocumentParser().parse(BytesIO(b"legacy"), "legacy.doc")

    assert parsed.parser == "antiword"
    assert parsed.page_numbering == "logical"
    assert [page.text_content for page in parsed.pages] == [
        "WPS No. WPS-003",
        "PQR No. PQR-003",
    ]


def test_legacy_doc_reports_missing_server_parser(monkeypatch) -> None:
    monkeypatch.setattr(document_parser_service.shutil, "which", lambda name: None)

    with pytest.raises(DocumentParseError, match="解析组件未安装"):
        DefaultDocumentParser().parse(BytesIO(b"legacy"), "legacy.doc")


def test_xlsx_roster_preserves_worksheets_and_rows() -> None:
    stream = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "焊工名册"
    sheet.append(["姓名", "焊工编号", "证书号", "有效期"])
    sheet.append(["张三", "W-001", "CERT-001", "2027-12-31"])
    workbook.create_sheet("续证记录").append(["李四", "W-002", "CERT-002"])
    workbook.save(stream)
    stream.seek(0)

    parsed = DefaultDocumentParser().parse(stream, "welders.xlsx")

    assert parsed.parser == "openpyxl"
    assert parsed.page_numbering == "worksheet"
    assert len(parsed.pages) == 2
    assert parsed.pages[0].metadata["sheet_name"] == "焊工名册"
    assert "张三\tW-001\tCERT-001" in parsed.pages[0].text_content
    assert parsed.pages[0].ocr_status == "not_required"


def test_docx_archive_expansion_limit_is_enforced(monkeypatch) -> None:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("PQR")
    document.save(stream)
    stream.seek(0)
    monkeypatch.setattr(document_parser_service, "MAX_DOCX_UNCOMPRESSED_BYTES", 1)

    with pytest.raises(DocumentParseError, match="解压后大小"):
        DefaultDocumentParser().parse(stream, "PQR.docx")


def test_pdf_page_limit_rejects_oversized_long_document(monkeypatch) -> None:
    monkeypatch.setattr(document_parser_service, "MAX_DOCUMENT_PAGES", 1)

    with pytest.raises(DocumentParseError, match="页数不能超过 1 页"):
        DefaultDocumentParser().parse(_text_pdf(), "long.pdf")


def test_page_text_limit_rejects_pathological_content(monkeypatch) -> None:
    monkeypatch.setattr(document_parser_service, "MAX_PAGE_TEXT_CHARS", 5)

    with pytest.raises(DocumentParseError, match="单页文本内容异常"):
        DefaultDocumentParser().parse(_text_pdf(), "large-text.pdf")


def test_batch_extraction_rejects_more_than_one_hundred_documents() -> None:
    with pytest.raises(ValidationError, match="at most 100 items"):
        BatchAIExtractionRequest(document_ids=[f"doc-{index}" for index in range(101)])


def test_hundred_page_text_pdf_parses_within_quality_budget() -> None:
    started = perf_counter()
    parsed = DefaultDocumentParser().parse(_many_page_pdf(100), "bulk-wps.pdf")

    assert len(parsed.pages) == 100
    assert perf_counter() - started < 10


def test_twenty_page_scanned_tiff_parses_within_quality_budget() -> None:
    stream = BytesIO()
    frames = [Image.new("1", (100, 100), 1) for _ in range(20)]
    frames[0].save(
        stream,
        format="TIFF",
        save_all=True,
        append_images=frames[1:],
    )
    stream.seek(0)
    started = perf_counter()

    parsed = DefaultDocumentParser().parse(stream, "bulk-scan.tiff")

    assert len(parsed.pages) == 20
    assert all(page.ocr_status == "pending" for page in parsed.pages)
    assert perf_counter() - started < 10
